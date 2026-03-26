"""
pipeline/resume_parser.py  — v2
────────────────────────────────
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Optional

import pdfplumber
import spacy
from docx import Document as DocxDocument
from loguru import logger

from config import settings


_NLP: Optional[spacy.Language] = None

def get_nlp() -> spacy.Language:
    global _NLP
    if _NLP is None:
        logger.info(f"Loading spaCy model: {settings.SPACY_MODEL}")
        _NLP = spacy.load(settings.SPACY_MODEL)
    return _NLP


TECH_SKILLS: set[str] = {
    "python", "java", "scala", "kotlin", "swift", "go", "golang", "rust",
    "c++", "c#", "ruby", "javascript", "typescript", "php", "r", "matlab",
    "julia", "dart", "bash", "shell", "perl", "groovy",
    "react", "vue", "angular", "svelte", "next.js", "nuxt", "html", "css",
    "sass", "tailwind", "bootstrap", "graphql", "webpack", "vite",
    "react.js", "angular.js", "vue.js",
    "fastapi", "django", "flask", "spring", "spring boot", "express",
    "express.js", "node.js", "nestjs", "laravel", "rails", "asp.net", "gin", "fiber",
    "docker", "kubernetes", "terraform", "ansible", "helm", "nginx",
    "apache", "jenkins", "ci/cd", "github actions", "gitlab ci",
    "aws", "azure", "gcp", "google cloud", "lambda", "ec2", "s3", "rds",
    "ecs", "eks", "cloudfront", "sqs", "sns", "eventbridge",
    "azure ml", "vertex ai", "kubeflow", "mlflow", "airflow",
    "machine learning", "deep learning", "nlp",
    "natural language processing", "computer vision",
    "llm", "transformers", "pytorch", "tensorflow", "keras",
    "scikit-learn", "xgboost", "lightgbm", "catboost",
    "spark", "pyspark", "hadoop", "kafka", "dbt",
    "mlops", "feast", "faiss",
    "hugging face", "langchain", "langgraph", "openai", "gemini",
    "rag", "retrieval augmented generation", "agentic", "prompt engineering",
    "generative ai", "llm evaluation", "vector database",
    "pandas", "numpy", "matplotlib", "seaborn", "plotly",
    "tableau", "power bi", "looker",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "cassandra", "dynamodb", "sqlite", "neo4j", "clickhouse",
    "snowflake", "bigquery", "databricks",
    "git", "github", "gitlab", "jira", "postman", "swagger",
    "streamlit", "gradio",
    "socket.io", "websockets", "grpc", "rest", "rest apis",
    "microservices", "etl", "data pipeline", "feature engineering",
    "new relic", "cloudwatch", "grafana", "prometheus",
    "aws solutions architect", "azure devops",
    "full stack", "devops", "agile", "scrum",
}

SKILL_SYNONYMS = {
    "ml": "machine learning", "ai": "artificial intelligence",
    "nlp": "natural language processing", "dl": "deep learning",
    "cv": "computer vision", "torch": "pytorch",
    "sklearn": "scikit-learn", "scikit learn": "scikit-learn",
    "postgres": "postgresql", "k8s": "kubernetes", "tf": "tensorflow",
    "js": "javascript", "ts": "typescript", "node": "node.js",
    "nodejs": "node.js", "react.js": "react", "reactjs": "react",
    "mongo": "mongodb", "llms": "llm",
    "large language models": "llm", "gen ai": "generative ai",
    "genai": "generative ai", "rag": "retrieval augmented generation",
}

def _norm(s: str) -> str:
    return SKILL_SYNONYMS.get(s.lower().strip(), s.lower().strip())

SKILL_BLACKLIST_PATTERNS = [
    r"^\d", r"%", r"\d{4}",
    r"jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec",
    r"present|current|till|since",
    r"^(hobbies|languages|education|experience|projects|"
    r"certifications?|awards?|skills?|summary|profile|"
    r"objective|references?|activities|co-curricular|"
    r"additional|links?|page\s*\d)$",
    r"travelling|cooking|listening|reading|gaming|photography",
    r"cgpa|gpa|marks|score|grade",
    r"intern|fresher|engineer\b",
    r"\band\b|\bthe\b|\bfor\b|\bwith\b|\busing\b",
    r"^\s*[•·\-–—]\s*$",
    r"[<>{}()\[\]\/\\]",
]

BLACKLIST_RE = [re.compile(p, re.IGNORECASE) for p in SKILL_BLACKLIST_PATTERNS]

def _is_valid_skill(s: str) -> bool:
    s = s.strip()
    if not s or len(s) < 2 or len(s) > 60:
        return False
    for pat in BLACKLIST_RE:
        if pat.search(s):
            return False
    return True


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_text_from_pdf_bytes(data: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            raw = page.extract_text()
            if raw:
                parts.append(raw)
    return "\n".join(parts)

def _extract_text_from_docx_bytes(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)

def extract_text_from_pdf(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_text_from_pdf_bytes(f.read())

def extract_text_from_docx(path: Path) -> str:
    with open(path, "rb") as f:
        return _extract_text_from_docx_bytes(f.read())

def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix in (".docx", ".doc"):
        return extract_text_from_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


# ── Field extractors ──────────────────────────────────────────────────────────

_EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+\-]+@[a-zA-Z0-9\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"(\+?\d[\d\s\-().]{7,14}\d)")

def extract_email(text: str) -> Optional[str]:
    m = _EMAIL_RE.search(text)
    return m.group(0).lower() if m else None

def extract_phone(text: str) -> Optional[str]:
    m = _PHONE_RE.search(text)
    return re.sub(r"\s+", " ", m.group(0)).strip() if m else None

def extract_experience_years(text: str) -> Optional[float]:
    patterns = [
        r"(?:over|more\s+than|around|approximately)?\s*(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)",
        r"experience\s*[:\-–]?\s*(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)",
        r"(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)\s+of\s+(?:professional|industry|work|overall|total)",
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return float(m.group(1))

    from datetime import datetime
    now = datetime.now()
    ranges = re.findall(
        r"(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?(\d{4})"
        r"\s*[-–—to]+\s*"
        r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}"
        r"|present|current|now|till\s+present|\d{4})",
        text, re.IGNORECASE,
    )
    total_months = 0
    seen = set()
    for start_str, end_str in ranges:
        key = (start_str, end_str.lower().strip())
        if key in seen:
            continue
        seen.add(key)
        try:
            s = int(start_str)
            end_clean = end_str.lower().strip()
            if any(w in end_clean for w in ("present", "current", "now", "till")):
                e, em = now.year, now.month
            else:
                e_match = re.search(r"\d{4}", end_str)
                e = int(e_match.group()) if e_match else now.year
                em = 12
            months = (e - s) * 12 + (em - 1)
            if 3 < months < 600:
                total_months += months
        except Exception:
            pass
    if total_months > 0:
        return round(total_months / 12, 1)
    return None

SKILL_SECTION_HEADERS_RE = re.compile(
    r"^(technical\s+skills?|skills?\s*(?:&\s*expertise)?|technologies|"
    r"tech(?:nical)?\s+stack|core\s+competencies?|expertise|proficiencies|"
    r"tools?\s*(?:&\s*technologies)?)",
    re.IGNORECASE,
)

NEXT_SECTION_RE = re.compile(
    r"^(experience|work\s+experience|employment|education|projects?|"
    r"certifications?|awards?|publications?|languages?|references?|"
    r"hobbies|interests?|summary|profile|objective|achievements?)",
    re.IGNORECASE,
)

def extract_skills(text: str) -> list[str]:
    found: set[str] = set()
    lower = text.lower()
    for skill in TECH_SKILLS:
        if re.search(r"\b" + re.escape(skill) + r"\b", lower):
            display = skill.title() if len(skill) <= 4 else skill.capitalize()
            found.add(display)
    lines = text.splitlines()
    in_skills = False
    for line in lines:
        stripped = line.strip()
        if SKILL_SECTION_HEADERS_RE.match(stripped):
            in_skills = True
            continue
        if in_skills:
            if not stripped or NEXT_SECTION_RE.match(stripped):
                in_skills = False
                continue
            parts = re.split(r"[,|•·;/\t]", stripped)
            for part in parts:
                clean = re.sub(r"^[\s•·\-–—:]+|[\s•·\-–—:]+$", "", part).strip()
                if _is_valid_skill(clean) and 2 <= len(clean) <= 50:
                    found.add(clean)
    final = [s for s in found if _is_valid_skill(s)]
    return sorted(set(final))

_TITLE_WORDS = {
    "senior", "lead", "principal", "staff", "junior", "associate",
    "engineer", "developer", "manager", "director", "analyst",
    "consultant", "architect", "scientist", "specialist", "officer",
    "intern", "fresher", "head", "vp", "cto", "ceo", "coo",
}

def _looks_like_name(text: str) -> bool:
    words = text.strip().split()
    if not (2 <= len(words) <= 4):
        return False
    if any(c.isdigit() for c in text):
        return False
    if not all(w[0].isupper() for w in words if w and w[0].isalpha()):
        return False
    if {w.lower() for w in words} & _TITLE_WORDS:
        return False
    return True

def extract_name(text: str, nlp: spacy.Language) -> Optional[str]:
    doc = nlp(text[:400])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            if 3 < len(name) < 60 and _looks_like_name(name):
                return name
    for line in text.splitlines()[:10]:
        stripped = line.strip()
        if not stripped or "@" in stripped or "http" in stripped:
            continue
        if re.search(r"\d{4}|\+\d|linkedin|github", stripped, re.IGNORECASE):
            continue
        if _looks_like_name(stripped):
            return stripped
    m = re.search(r"(?:full\s+)?name\s*[:\-]\s*([A-Z][a-z]+(?: [A-Z][a-z]+)+)", text, re.I)
    if m:
        return m.group(1).strip()
    email = extract_email(text)
    if email:
        local = email.split("@")[0]
        parts = re.split(r"[._\-]", local)
        parts = [p.capitalize() for p in parts if p.isalpha() and len(p) > 1]
        if 2 <= len(parts) <= 3:
            return " ".join(parts)
    return None

_DEGREE_RE = re.compile(
    r"\b(b\.?tech|m\.?tech|b\.?e\b|m\.?e\b|bachelor|master|phd|ph\.d|"
    r"mba|m\.?sc|b\.?sc|b\.?com|diploma|graduate|postgraduate|"
    r"doctor(?:ate)?|associate|b\.?s\b|m\.?s\b)\b",
    re.IGNORECASE,
)

def extract_education(text: str) -> Optional[str]:
    edu_lines: list[str] = []
    for line in text.splitlines():
        if _DEGREE_RE.search(line):
            clean = line.strip()
            if 5 < len(clean) < 200:
                edu_lines.append(clean)
    return " | ".join(edu_lines[:3]) if edu_lines else None

COMPANY_SUFFIXES = re.compile(
    r"\b(pvt\.?\s*ltd\.?|ltd\.?|llc|inc\.?|corp\.?|corporation|"
    r"technologies|technology|solutions|systems|services|consulting|"
    r"consultancy|analytics|labs?|studio|ventures|group|global|"
    r"management|enterprises?|networks?)\b",
    re.IGNORECASE,
)

EXPERIENCE_COMPANY_RE = re.compile(
    r"^([A-Z][A-Za-z\s&,.()\-/]+?)\s+"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|[A-Z][a-z]+)?\s*"
    r"(?:19|20)\d{2}",
    re.MULTILINE,
)

def extract_companies(text: str, nlp: spacy.Language) -> list[str]:
    companies: set[str] = set()
    doc = nlp(text[:4000])
    for ent in doc.ents:
        if ent.label_ == "ORG":
            name = ent.text.strip()
            if (COMPANY_SUFFIXES.search(name) or len(name.split()) >= 2) and len(name) < 80:
                companies.add(name)
    for m in EXPERIENCE_COMPANY_RE.finditer(text):
        candidate = m.group(1).strip().rstrip(",")
        if 3 < len(candidate) < 60 and not candidate.isupper():
            companies.add(candidate)
    NON_COMPANY = {"bachelor", "master", "university", "college", "school", "institute", "education", "skills", "experience", "projects"}
    result = [c for c in companies if not any(w.lower() in NON_COMPANY for w in c.split()) and len(c) > 3]
    return list(dict.fromkeys(result))[:8]

CERT_SECTION_RE = re.compile(
    r"^(certifications?|courses?\s*(?:&\s*certifications?)?|"
    r"licenses?\s*(?:&\s*certifications?)?|professional\s+development|"
    r"training\s*(?:&\s*certifications?)?)",
    re.IGNORECASE,
)

CERT_END_RE = re.compile(
    r"^(awards?|education|projects?|experience|work|employment|"
    r"skills?|languages?|references?|publications?|hobbies|additional|links?)",
    re.IGNORECASE,
)

CERT_KEYWORD_RE = re.compile(
    r"\b(certified|certification|certificate|aws\s|azure\s|google\s+cloud|"
    r"gcp\s|microsoft\s|databricks|coursera|udemy|edx|stanford|deeplearning|"
    r"hackerrank|leetcode|udacity|pmp|cka|cks|scrum|cissp|comptia)\b",
    re.IGNORECASE,
)

def extract_certifications(text: str) -> list[str]:
    certs: list[str] = []
    lines = text.splitlines()
    in_cert_section = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if CERT_SECTION_RE.match(stripped):
            in_cert_section = True
            continue
        if in_cert_section:
            if CERT_END_RE.match(stripped):
                in_cert_section = False
                continue
            clean = re.sub(r"^[\s•·\-–—▪◦\uf0b7]+", "", stripped).strip()
            if 5 < len(clean) < 200:
                certs.append(clean)
    if not certs:
        for line in lines:
            stripped = line.strip()
            if CERT_KEYWORD_RE.search(stripped):
                clean = re.sub(r"^[\s•·\-–—▪◦\uf0b7]+", "", stripped).strip()
                if 5 < len(clean) < 200:
                    if not re.search(r"(experience|developed|built|managed|led|designed)", clean, re.I):
                        certs.append(clean)
    return list(dict.fromkeys(certs))[:12]

PROJECT_SECTION_RE = re.compile(
    r"^((?:key\s+|notable\s+|personal\s+|major\s+)?projects?|"
    r"academic\s+projects?|portfolio)",
    re.IGNORECASE,
)

PROJECT_END_RE = re.compile(
    r"^(education|certifications?|skills?|awards?|languages?|"
    r"references?|hobbies|additional|links?|publications?)",
    re.IGNORECASE,
)

def extract_projects(text: str) -> list[str]:
    projects: list[str] = []
    lines = text.splitlines()
    in_projects = False
    current_project: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if PROJECT_SECTION_RE.match(stripped):
            in_projects = True
            continue
        if in_projects:
            if PROJECT_END_RE.match(stripped):
                if current_project:
                    projects.append(" ".join(current_project[:2]))
                break
            clean = re.sub(r"^[\s•·\-–—▪◦\uf0b7]+", "", stripped).strip()
            if "|" in clean and len(clean) < 200:
                if current_project:
                    projects.append(current_project[0])
                current_project = [clean.split("|")[0].strip()]
            elif not clean.startswith("–") and not clean.startswith("-") and 5 < len(clean) < 100:
                if len(current_project) == 0:
                    current_project = [clean]
                elif len(current_project) == 1:
                    current_project.append(clean)
            if len(projects) >= 10:
                break
    if current_project and len(projects) < 10:
        projects.append(current_project[0])
    return list(dict.fromkeys(projects))[:10]


# ── Shared internal parser ────────────────────────────────────────────────────

def _parse_raw_text(raw_text: str, file_name: str, file_path: str) -> dict:
    raw_text = re.sub(r"[\uf0b7\uf0a7\uf0d8\uf076]", "•", raw_text)
    raw_text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", raw_text)
    nlp = get_nlp()
    result: dict = {
        "name":               extract_name(raw_text, nlp),
        "email":              extract_email(raw_text),
        "phone":              extract_phone(raw_text),
        "skills":             extract_skills(raw_text),
        "education":          extract_education(raw_text),
        "experience_years":   extract_experience_years(raw_text),
        "previous_companies": extract_companies(raw_text, nlp),
        "certifications":     extract_certifications(raw_text),
        "projects":           extract_projects(raw_text),
        "raw_text":           raw_text,
        "file_name":          file_name,
        "file_path":          file_path,
    }
    logger.debug(
        f"Parsed {file_name}: name={result['name']}, "
        f"skills={len(result['skills'])}, exp={result['experience_years']}yrs"
    )
    return result


# ── Public API ────────────────────────────────────────────────────────────────

def parse_resume_bytes(content: bytes, suffix: str, filename: str = "") -> dict:
    """Parse resume from raw bytes — for uploaded/cloud files."""
    fname = filename or f"uploaded{suffix}"
    try:
        if suffix == ".pdf":
            raw_text = _extract_text_from_pdf_bytes(content)
        elif suffix in (".docx", ".doc"):
            raw_text = _extract_text_from_docx_bytes(content)
        else:
            raw_text = content.decode("utf-8", errors="ignore")
    except Exception as exc:
        logger.error(f"Text extraction failed for {fname}: {exc}")
        return {"raw_text": "", "file_name": fname, "file_path": fname}
    return _parse_raw_text(raw_text, fname, fname)


def parse_resume(path: Path) -> dict:
    """Parse resume from file path — for local files."""
    try:
        raw_text = extract_text(path)
    except Exception as exc:
        logger.error(f"Text extraction failed for {path.name}: {exc}")
        return {"raw_text": "", "file_name": path.name, "file_path": str(path)}
    return _parse_raw_text(raw_text, path.name, str(path))
"""
nlp/resume_parser.py  — v3 (Production-grade ATS parser)
──────────────────────────────────────────────────────────
Core philosophy:
  1. Parse the resume into NAMED SECTIONS first (boundary detection)
  2. Extract each field ONLY from the correct section
  3. Apply validation + post-processing to remove noise/misclassified entries
  4. Normalize entities (company names, degree titles, cert formats)

Section types recognised:
  CONTACT | SUMMARY | SKILLS | EXPERIENCE | EDUCATION | CERTIFICATIONS |
  PROJECTS | AWARDS | LANGUAGES | REFERENCES | OTHER

Fields extracted:
  name, email, phone, skills, experience_years,
  education (list of structured dicts), previous_companies (list),
  certifications (list), projects (list)
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pdfplumber
import spacy
from docx import Document as DocxDocument
from loguru import logger

from config import settings


# ── spaCy singleton ───────────────────────────────────────────────────────────

_NLP: Optional[spacy.Language] = None


def get_nlp() -> spacy.Language:
    global _NLP
    if _NLP is None:
        logger.info(f"Loading spaCy model: {settings.SPACY_MODEL}")
        try:
            _NLP = spacy.load(settings.SPACY_MODEL)
        except OSError:
            raise OSError(
                f"spaCy model '{settings.SPACY_MODEL}' not found. "
                "Run: python -m spacy download en_core_web_sm"
            )
    return _NLP


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 1 — TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(path: Path) -> str:
    parts: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                raw = page.extract_text(x_tolerance=2, y_tolerance=2)
                if raw:
                    parts.append(raw)
    except Exception as e:
        logger.warning(f"pdfplumber failed on {path.name}: {e}")
    return "\n".join(parts)


def extract_text_from_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    lines.append(t)
    return "\n".join(lines)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        raw = extract_text_from_pdf(path)
    elif suffix in (".docx", ".doc"):
        raw = extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    raw = re.sub(r"[\uf0b7\uf0a7\uf0d8\uf076\u2022\u25cf\u25aa\u2713]", "•", raw)
    raw = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", raw)
    raw = re.sub(r"\r\n|\r", "\n", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw)
    return raw.strip()


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 2 — SECTION BOUNDARY DETECTION
# ═══════════════════════════════════════════════════════════════════════════════

class SectionType:
    CONTACT        = "CONTACT"
    SUMMARY        = "SUMMARY"
    SKILLS         = "SKILLS"
    EXPERIENCE     = "EXPERIENCE"
    EDUCATION      = "EDUCATION"
    CERTIFICATIONS = "CERTIFICATIONS"
    PROJECTS       = "PROJECTS"
    AWARDS         = "AWARDS"
    LANGUAGES      = "LANGUAGES"
    REFERENCES     = "REFERENCES"
    OTHER          = "OTHER"


_SECTION_PATTERNS: list[tuple[str, str]] = [
    (SectionType.EDUCATION,
     r"^(education(?:al)?\s*(?:background|qualifications?|details?)?|"
     r"academic\s*(?:background|qualifications?|history|profile)?|"
     r"qualifications?|degrees?)\s*:?\s*$"),

    (SectionType.EXPERIENCE,
     r"^((?:professional|work|career|employment|job|industry)\s*"
     r"(?:experience|history|background|summary)?|"
     r"experience|work\s*history|employment\s*(?:history|details?)?|"
     r"professional\s*background|career\s*(?:history|summary)?)\s*:?\s*$"),

    (SectionType.CERTIFICATIONS,
     r"^(certifications?\s*(?:&\s*licen[sc]es?)?|"
     r"licen[sc]es?\s*(?:&\s*certifications?)?|"
     r"professional\s*certifications?|"
     r"courses?\s*(?:&\s*certifications?)?|"
     r"training\s*(?:&\s*certifications?)?|"
     r"credentials?)\s*:?\s*$"),

    (SectionType.SKILLS,
     r"^(technical\s*skills?|skills?\s*(?:&\s*expertise|summary)?|"
     r"core\s*(?:competencies?|skills?)|competencies?|"
     r"technologies|tech(?:nical)?\s*stack|expertise|"
     r"tools?\s*(?:&\s*technologies?)?|proficiencies?|"
     r"key\s*skills?|functional\s*skills?)\s*:?\s*$"),

    (SectionType.PROJECTS,
     r"^((?:key|notable|major|personal|academic|significant)\s*"
     r"projects?|projects?|portfolio|key\s*achievements?)\s*:?\s*$"),

    (SectionType.SUMMARY,
     r"^((?:professional\s*)?summary|profile|objective|"
     r"career\s*objective|about\s*(?:me)?|overview|"
     r"executive\s*summary|professional\s*profile)\s*:?\s*$"),

    (SectionType.AWARDS,
     r"^(awards?\s*(?:&\s*(?:honors?|achievements?))?|"
     r"honors?|achievements?|accomplishments?|recognition)\s*:?\s*$"),

    (SectionType.LANGUAGES,
     r"^(languages?|language\s*skills?|linguistic\s*skills?)\s*:?\s*$"),

    (SectionType.REFERENCES,
     r"^(references?|referees?)\s*:?\s*$"),
]

_COMPILED_SECTIONS = [
    (stype, re.compile(pat, re.IGNORECASE))
    for stype, pat in _SECTION_PATTERNS
]


def _classify_line_as_section(line: str) -> Optional[str]:
    stripped = line.strip()
    if not stripped or len(stripped) > 80:
        return None
    alpha_ratio = sum(c.isalpha() or c.isspace() for c in stripped) / len(stripped)
    if alpha_ratio < 0.70:
        return None
    for stype, pat in _COMPILED_SECTIONS:
        if pat.match(stripped):
            return stype
    return None


@dataclass
class Section:
    type: str
    lines: list[str] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(self.lines)

    def add(self, line: str) -> None:
        self.lines.append(line)


def parse_sections(raw_text: str) -> dict[str, list[Section]]:
    sections: dict[str, list[Section]] = {}
    current_section = Section(type=SectionType.OTHER)

    for line in raw_text.splitlines():
        stype = _classify_line_as_section(line)
        if stype:
            if current_section.lines:
                sections.setdefault(current_section.type, []).append(current_section)
            current_section = Section(type=stype)
        else:
            current_section.add(line)

    if current_section.lines:
        sections.setdefault(current_section.type, []).append(current_section)

    return sections


def _get_section_text(sections: dict, stype: str) -> str:
    parts = sections.get(stype, [])
    return "\n".join(s.text for s in parts)


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 3 — CONTACT FIELDS
# ═══════════════════════════════════════════════════════════════════════════════

_EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+\-]+@[a-zA-Z0-9\-]+\.[a-zA-Z]{2,}")
_PHONE_RE = re.compile(r"(\+?(?:\d[\s\-.]?){7,14}\d)")

_TITLE_WORDS = {
    "senior", "lead", "principal", "staff", "junior", "associate", "mid",
    "engineer", "developer", "manager", "director", "analyst", "vp",
    "consultant", "architect", "scientist", "specialist", "officer",
    "intern", "head", "cto", "ceo", "coo", "cfo", "svp", "avp",
    "president", "founder", "co-founder", "partner", "executive",
    "programmer", "designer", "tester", "qa", "devops", "sre",
    "fullstack", "frontend", "backend", "mobile", "cloud",
}


def extract_email(text: str) -> Optional[str]:
    m = _EMAIL_RE.search(text)
    return m.group(0).lower() if m else None


def extract_phone(text: str) -> Optional[str]:
    for line in text.splitlines()[:30]:
        m = _PHONE_RE.search(line)
        if m:
            phone = re.sub(r"[\s\-.]", " ", m.group(0)).strip()
            digits = re.sub(r"\D", "", phone)
            if 7 <= len(digits) <= 15:
                return phone
    return None


def _looks_like_name(text: str) -> bool:
    words = text.strip().split()
    if not (2 <= len(words) <= 5):
        return False
    if any(c.isdigit() for c in text):
        return False
    if not all(w[0].isupper() for w in words if w and w[0].isalpha()):
        return False
    lwords = {w.lower().rstrip(".,") for w in words}
    if lwords & _TITLE_WORDS:
        return False
    if re.search(
        r"(street|road|avenue|lane|nagar|city|state|district|"
        r"india|usa|uk|dubai|singapore|resume|curriculum|cv\b)",
        text, re.IGNORECASE
    ):
        return False
    return True


def _normalize_name(name: str) -> str:
    name = re.sub(r"\s+", " ", name).strip()
    return " ".join(w.capitalize() for w in name.split())


def extract_name(text: str, nlp: spacy.Language) -> Optional[str]:
    header = text[:600]

    # Strategy 1: spaCy PERSON
    doc = nlp(header)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            if 3 < len(name) < 60 and _looks_like_name(name):
                return _normalize_name(name)

    # Strategy 2: First non-metadata line
    for line in text.splitlines()[:15]:
        stripped = line.strip()
        if not stripped:
            continue
        if re.search(
            r"[@|http|linkedin|github|\d{4}|\+\d|curriculum|resume|cv\b]",
            stripped, re.IGNORECASE
        ):
            continue
        if _looks_like_name(stripped):
            return _normalize_name(stripped)

    # Strategy 3: "Name:" label
    m = re.search(
        r"(?:full\s+)?name\s*[:\-]\s*([A-Z][a-zA-Z]+(?: [A-Z][a-zA-Z]+)+)",
        header, re.IGNORECASE
    )
    if m:
        return _normalize_name(m.group(1).strip())

    # Strategy 4: Derive from email
    email = extract_email(text)
    if email:
        local = email.split("@")[0]
        parts = re.split(r"[._\-]", local)
        parts = [p.capitalize() for p in parts if p.isalpha() and len(p) > 1]
        if 2 <= len(parts) <= 3:
            return " ".join(parts)

    return None


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 4 — EDUCATION (strict: degrees + institutions only)
# ═══════════════════════════════════════════════════════════════════════════════

_DEGREE_ALIASES: dict[str, str] = {
    "btech": "B.Tech", "b.tech": "B.Tech", "b tech": "B.Tech",
    "mtech": "M.Tech", "m.tech": "M.Tech", "m tech": "M.Tech",
    "be": "B.E", "b.e": "B.E", "me": "M.E", "m.e": "M.E",
    "bsc": "B.Sc", "b.sc": "B.Sc", "b sc": "B.Sc",
    "msc": "M.Sc", "m.sc": "M.Sc", "m sc": "M.Sc",
    "bca": "BCA", "mca": "MCA", "bcom": "B.Com", "b.com": "B.Com",
    "mba": "MBA", "phd": "Ph.D", "ph.d": "Ph.D", "doctorate": "Ph.D",
    "bachelor": "Bachelor's", "bachelor's": "Bachelor's",
    "master": "Master's", "master's": "Master's",
    "diploma": "Diploma", "associate": "Associate's",
    "bs": "B.S", "b.s": "B.S", "ms": "M.S", "m.s": "M.S",
    "postgraduate": "Postgraduate", "graduate": "Graduate",
    "hsc": "HSC", "h.s.c": "HSC", "sslc": "SSLC", "s.s.l.c": "SSLC",
    "10th": "Class X", "12th": "Class XII",
    "higher secondary": "Higher Secondary",
    "secondary school": "Secondary School",
}

_DEGREE_RE = re.compile(
    r"\b(b\.?tech|m\.?tech|b\.?e\b|m\.?e\b|b\.?sc|m\.?sc|bca|mca|"
    r"b\.?com|mba|ph\.?d|doctorate|bachelor(?:'s)?|master(?:'s)?|"
    r"diploma|associate(?:'s)?|b\.?s\b|m\.?s\b|postgraduate|graduate|"
    r"hsc|sslc|h\.s\.c|s\.s\.l\.c|higher\s+secondary|secondary\s+school|"
    r"10th|12th|class\s*x{1,2}|std\s*\d+)\b",
    re.IGNORECASE,
)

_EDU_NOISE_RE = re.compile(
    r"\b(present|current|till|ongoing|responsibilities?|achieved|managed|"
    r"developed|implemented|led|designed|built|created|worked|handled|"
    r"responsible|ensure|support|provide|maintain|assist|coordinate|"
    r"prepare|execute|deliver)\b",
    re.IGNORECASE,
)

_INSTITUTION_RE = re.compile(
    r"\b(university|college|institute|school|academy|polytechnic|"
    r"iit|iim|nit|bits|vit|srm|anna\s+university|deemed)\b",
    re.IGNORECASE,
)

_YEAR_RE = re.compile(r"\b(19|20)\d{2}\b")


@dataclass
class EducationEntry:
    degree: str
    institution: Optional[str] = None
    year: Optional[str] = None
    field_of_study: Optional[str] = None

    def display(self) -> str:
        parts = [self.degree]
        if self.field_of_study:
            parts[0] += f" in {self.field_of_study}"
        if self.institution:
            parts.append(self.institution)
        if self.year:
            parts.append(f"({self.year})")
        return " | ".join(parts)


def _normalize_degree(raw: str) -> str:
    key = raw.lower().strip()
    return _DEGREE_ALIASES.get(key, raw.strip())


def _extract_year_from_line(line: str) -> Optional[str]:
    years = _YEAR_RE.findall(line)
    return max(years) if years else None


def _is_education_noise(line: str) -> bool:
    if _EDU_NOISE_RE.search(line):
        return True
    if re.match(r"^[•\-\*]\s+", line):
        if re.search(
            r"\b(developed|managed|led|designed|built|implemented|"
            r"achieved|delivered|maintained)\b", line, re.IGNORECASE
        ):
            return True
    return False


def extract_education(sections: dict) -> list[dict]:
    edu_text = _get_section_text(sections, SectionType.EDUCATION)
    if not edu_text.strip():
        edu_text = _get_section_text(sections, SectionType.OTHER)

    entries: list[EducationEntry] = []
    lines = edu_text.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or _is_education_noise(line):
            i += 1
            continue

        degree_match = _DEGREE_RE.search(line)
        if not degree_match:
            i += 1
            continue

        raw_degree = degree_match.group(0)
        degree = _normalize_degree(raw_degree)

        post_degree = line[degree_match.end():].strip()
        post_degree = re.sub(r"^[\s\-–,in]+", "", post_degree)
        field = None
        if post_degree and len(post_degree) < 60 and not _YEAR_RE.match(post_degree):
            candidate_field = re.sub(
                r"\b(from|at|in|the|and|or|of)\b", "", post_degree, flags=re.I
            ).strip()
            if candidate_field and not _INSTITUTION_RE.search(candidate_field):
                field = candidate_field[:50]

        year = _extract_year_from_line(line)

        institution = None
        if _INSTITUTION_RE.search(line):
            inst_match = re.search(
                r"([A-Z][A-Za-z\s&.,()'\-]+(?:"
                r"university|college|institute|school|academy|polytechnic|"
                r"iit|iim|nit|bits|vit|srm))",
                line, re.IGNORECASE
            )
            if inst_match:
                institution = _normalize_institution(inst_match.group(0).strip())
        else:
            for j in range(i + 1, min(i + 4, len(lines))):
                next_line = lines[j].strip()
                if not next_line:
                    continue
                if _INSTITUTION_RE.search(next_line) or (
                    len(next_line.split()) >= 2
                    and next_line[0].isupper()
                    and not _DEGREE_RE.search(next_line)
                    and not _is_education_noise(next_line)
                ):
                    institution = _normalize_institution(next_line)
                    if not year:
                        year = _extract_year_from_line(next_line)
                    break

        entry = EducationEntry(
            degree=degree, institution=institution,
            year=year, field_of_study=field,
        )

        if not any(
            e.degree == entry.degree and e.institution == entry.institution
            for e in entries
        ):
            entries.append(entry)

        i += 1

    return [
        {
            "degree": e.degree,
            "institution": e.institution,
            "year": e.year,
            "field_of_study": e.field_of_study,
            "display": e.display(),
        }
        for e in entries[:6]
    ]


def _normalize_institution(name: str) -> str:
    name = re.sub(r"\s+", " ", name).strip()
    name = re.sub(r"\s*[,\-–]\s*$", "", name)
    return name[:120]


def education_display(edu_list: list[dict]) -> Optional[str]:
    if not edu_list:
        return None
    return " | ".join(e["display"] for e in edu_list)


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 5 — WORK EXPERIENCE / COMPANIES
# ═══════════════════════════════════════════════════════════════════════════════

_COMPANY_NOISE_WORDS = {
    "university", "college", "school", "institute", "academy", "polytechnic",
    "iit", "iim", "nit", "education", "skills", "projects", "certifications",
    "awards", "languages", "references", "summary", "profile",
    "objective", "hobbies", "interests", "activities", "team", "department",
}

_COMPANY_ROLE_INDICATORS = re.compile(
    r"^(software|senior|lead|principal|manager|director|engineer|"
    r"developer|analyst|architect|consultant|specialist|officer|"
    r"associate|intern|executive|head|vp|cto|ceo|president|founder)\b",
    re.IGNORECASE,
)

_COMPANY_SUFFIXES_RE = re.compile(
    r"\b(pvt\.?\s*ltd\.?|ltd\.?|llc|inc\.?|corp\.?|corporation|"
    r"technologies|technology|solutions|systems|services|consulting|"
    r"consultancy|analytics|labs?|studio|ventures|group|global|"
    r"management|enterprises?|networks?|digital|software|infotech|"
    r"infosystems?|techno(?:logies)?|co\b)\b",
    re.IGNORECASE,
)

_DATE_RANGE_RE = re.compile(
    r"(?:"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?"
    r"(?:\s+\d{4})?"
    r"|(?:19|20)\d{2}"
    r")"
    r"\s*[-–—to]+\s*"
    r"(?:"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?"
    r"(?:\s+\d{4})?"
    r"|(?:19|20)\d{2}"
    r"|present|current|now|till\s+(?:present|date)|to\s+date"
    r")",
    re.IGNORECASE,
)


def _is_valid_company(name: str) -> bool:
    if not name or len(name) < 3 or len(name) > 100:
        return False
    words = set(re.split(r"\s+", name.lower()))
    if words & _COMPANY_NOISE_WORDS:
        return False
    if _DEGREE_RE.search(name):
        return False
    words_list = name.split()
    if len(words_list) <= 3 and _COMPANY_ROLE_INDICATORS.match(name):
        return False
    if re.match(r"^\d+$", name):
        return False
    return True


def _normalize_company(name: str) -> str:
    name = re.sub(r"\s+", " ", name).strip()
    name = re.sub(r"[\s,\.\-–]+$", "", name)
    name = re.sub(r"\bPvt\b", "Pvt.", name)
    name = re.sub(r"\bLtd\b", "Ltd.", name)
    return name[:100]


def extract_companies(sections: dict, nlp: spacy.Language) -> list[str]:
    exp_text = _get_section_text(sections, SectionType.EXPERIENCE)
    if not exp_text.strip():
        exp_text = _get_section_text(sections, SectionType.OTHER)

    companies: list[str] = []
    seen: set[str] = set()
    lines = exp_text.splitlines()

    # Strategy A: Lines with date ranges
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or not _DATE_RANGE_RE.search(stripped):
            continue
        candidates = [stripped, lines[i - 1].strip() if i > 0 else ""]
        for cand in candidates:
            cand = _DATE_RANGE_RE.sub("", cand).strip()
            cand = re.sub(r"[|\-–•:,]+$", "", cand).strip()
            if _is_valid_company(cand) and len(cand) > 3:
                norm = _normalize_company(cand)
                if norm.lower() not in seen:
                    companies.append(norm)
                    seen.add(norm.lower())

    # Strategy B: spaCy ORG entities
    doc = nlp(exp_text[:5000])
    for ent in doc.ents:
        if ent.label_ == "ORG":
            name = ent.text.strip()
            if not _is_valid_company(name):
                continue
            if _COMPANY_SUFFIXES_RE.search(name) or len(name.split()) >= 2:
                norm = _normalize_company(name)
                if norm.lower() not in seen:
                    companies.append(norm)
                    seen.add(norm.lower())

    return [c for c in companies if _is_valid_company(c)][:8]


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 6 — EXPERIENCE YEARS (with overlap merging)
# ═══════════════════════════════════════════════════════════════════════════════

def extract_experience_years(text: str) -> Optional[float]:
    # Explicit statement patterns
    explicit_patterns = [
        r"(?:over|more\s+than|around|approximately|nearly|about)?\s*"
        r"(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)\s*(?:of\s+)?"
        r"(?:professional\s+|industry\s+|overall\s+|total\s+|work\s+)?"
        r"(?:experience|exp|expertise)",
        r"experience\s*(?:of\s+)?(?:over\s+|more\s+than\s+)?(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)",
        r"(\d+\.?\d*)\s*\+?\s*(?:years?|yrs?)\s+of\s+(?:professional|industry|work|overall|total)",
    ]
    for pat in explicit_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            val = float(m.group(1))
            if 0 < val < 55:
                return val

    # Date-range accumulation with overlap merging
    from datetime import datetime
    now = datetime.now()

    _MONTH_MAP = {
        "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
        "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    }

    range_re = re.compile(
        r"(?:(?P<sm>jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)?"
        r"(?P<sy>(?:19|20)\d{2})"
        r"\s*[-–—to]+\s*"
        r"(?:"
        r"(?:(?P<em>jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+)"
        r"(?P<ey>(?:19|20)\d{2})"
        r"|(?P<now>present|current|now|till\s+(?:present|date)|to\s+date)"
        r"|(?P<ey2>(?:19|20)\d{2})"
        r")",
        re.IGNORECASE,
    )

    intervals: list[tuple[int, int]] = []
    for m in range_re.finditer(text):
        sy = int(m.group("sy"))
        sm = _MONTH_MAP.get((m.group("sm") or "jan")[:3].lower(), 1)
        if m.group("now"):
            ey, em = now.year, now.month
        elif m.group("ey"):
            ey = int(m.group("ey"))
            em = _MONTH_MAP.get((m.group("em") or "dec")[:3].lower(), 12)
        elif m.group("ey2"):
            ey = int(m.group("ey2"))
            em = 12
        else:
            continue
        if not (1970 <= sy <= now.year and 1970 <= ey <= now.year + 1):
            continue
        if sy > ey:
            continue
        intervals.append((sy * 12 + sm, ey * 12 + em))

    if not intervals:
        return None

    # Merge overlapping
    intervals.sort()
    merged: list[tuple[int, int]] = []
    for s, e in intervals:
        if merged and s <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], e))
        else:
            merged.append((s, e))

    total_months = sum(e - s for s, e in merged)
    if total_months <= 0:
        return None
    years = round(total_months / 12, 1)
    return years if 0.1 < years < 55 else None


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 7 — CERTIFICATIONS (strict: verifiable credentials only)
# ═══════════════════════════════════════════════════════════════════════════════

_CERT_ISSUERS = re.compile(
    r"\b(aws|amazon\s+web\s+services|azure|microsoft|google\s+cloud|gcp|"
    r"coursera|udemy|edx|stanford|mit|deeplearning\.?ai|deeplearning|"
    r"hackerrank|udacity|pluralsight|linkedin\s+learning|"
    r"databricks|snowflake|salesforce|oracle|cisco|comptia|pmi|"
    r"scrum\s*alliance|scrum\.org|isc2|isaca|ec-council|"
    r"red\s+hat|pmp|cka|cks|ckad|cissp|cism|cisa|"
    r"itil|prince2|six\s+sigma|pmbok|safe|togaf|"
    r"tensorflow|pytorch|nvidia|ibm|sap|tableau|"
    r"nptel|swayam|infosys\s+springboard|nasscom)\b",
    re.IGNORECASE,
)

_CERT_KEYWORDS = re.compile(
    r"\b(certified|certification|certificate|credential|license|licence|"
    r"accredited|completion|issued\s+by|awarded\s+by|"
    r"professional\s+development|course\s+completion)\b",
    re.IGNORECASE,
)

_CERT_ANTI_PATTERNS = re.compile(
    r"\b(b\.?tech|m\.?tech|b\.?e|m\.?e|b\.?sc|m\.?sc|mba|phd|ph\.d|"
    r"bachelor|master|diploma|degree|university|college|institute|"
    r"school\s+of|graduated|graduation|admission|"
    r"sslc|hsc|10th|12th|class\s*x|std\s*\d+|"
    r"affiliated|accredited\s+by|approved\s+by|autonomous)\b",
    re.IGNORECASE,
)

_CERT_EXPERIENCE_RE = re.compile(
    r"\b(developed|built|managed|led|designed|implemented|delivered|"
    r"worked|maintained|achieved|responsible|handled|coordinated|"
    r"created|supported|performed)\b",
    re.IGNORECASE,
)


def _is_valid_cert(line: str) -> bool:
    if not line or len(line) < 5 or len(line) > 300:
        return False
    if _CERT_ANTI_PATTERNS.search(line):
        return False
    if _CERT_EXPERIENCE_RE.search(line):
        return False
    return bool(_CERT_KEYWORDS.search(line)) or bool(_CERT_ISSUERS.search(line))


def _normalize_cert(cert: str) -> str:
    cert = re.sub(r"^[\s•·\-–—▪◦\uf0b7\*]+", "", cert).strip()
    cert = re.sub(r"\s+", " ", cert)
    cert = re.sub(r"\s*[|:,\-–]\s*$", "", cert)
    return cert[:250]


def extract_certifications(sections: dict) -> list[str]:
    cert_text = _get_section_text(sections, SectionType.CERTIFICATIONS)
    certs: list[str] = []
    seen: set[str] = set()

    if cert_text.strip():
        for line in cert_text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            clean = _normalize_cert(stripped)
            if not clean or len(clean) < 5:
                continue
            if _CERT_ANTI_PATTERNS.search(clean):
                continue
            if _CERT_EXPERIENCE_RE.search(clean):
                continue
            key = clean.lower()[:60]
            if key not in seen:
                certs.append(clean)
                seen.add(key)

    # Fallback scan — strict
    if not certs:
        scan_text = (
            _get_section_text(sections, SectionType.OTHER)
            + "\n"
            + _get_section_text(sections, SectionType.EXPERIENCE)
        )
        for line in scan_text.splitlines():
            stripped = line.strip()
            if not stripped or len(stripped) < 10:
                continue
            if _is_valid_cert(stripped):
                clean = _normalize_cert(stripped)
                key = clean.lower()[:60]
                if key not in seen:
                    certs.append(clean)
                    seen.add(key)

    # Remove header leakage
    final = [
        c for c in certs
        if not re.match(r"^(certifications?|licenses?|courses?)$", c, re.IGNORECASE)
    ]
    return list(dict.fromkeys(final))[:12]


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 8 — SKILLS
# ═══════════════════════════════════════════════════════════════════════════════

TECH_SKILLS: set[str] = {
    "python", "java", "scala", "kotlin", "swift", "go", "golang", "rust",
    "c++", "c#", "ruby", "javascript", "typescript", "php", "r", "matlab",
    "julia", "dart", "bash", "shell", "perl", "groovy", "cobol",
    "react", "vue", "angular", "svelte", "next.js", "nuxt", "html", "css",
    "sass", "tailwind", "bootstrap", "graphql", "webpack", "vite",
    "react.js", "angular.js", "vue.js", "redux", "jquery",
    "fastapi", "django", "flask", "spring", "spring boot", "express",
    "express.js", "node.js", "nestjs", "laravel", "rails", "asp.net",
    "docker", "kubernetes", "terraform", "ansible", "helm", "nginx",
    "apache", "jenkins", "ci/cd", "github actions", "gitlab ci",
    "aws", "azure", "gcp", "google cloud", "lambda", "ec2", "s3", "rds",
    "ecs", "eks", "azure ml", "vertex ai", "kubeflow", "mlflow", "airflow",
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "llm", "transformers", "pytorch", "tensorflow", "keras",
    "scikit-learn", "xgboost", "lightgbm", "catboost",
    "spark", "pyspark", "hadoop", "kafka", "dbt", "mlops", "faiss",
    "hugging face", "langchain", "langgraph", "openai", "gemini", "rag",
    "retrieval augmented generation", "prompt engineering", "generative ai",
    "vector database", "pandas", "numpy", "matplotlib", "seaborn", "plotly",
    "tableau", "power bi", "looker", "qlik",
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
    "cassandra", "dynamodb", "sqlite", "neo4j", "clickhouse",
    "snowflake", "bigquery", "databricks", "oracle", "sql server",
    "git", "github", "gitlab", "jira", "confluence", "postman", "swagger",
    "streamlit", "gradio", "grpc", "rest", "rest apis",
    "microservices", "etl", "data pipeline", "feature engineering",
    "grafana", "prometheus", "kibana", "datadog",
    "oracle erp", "oracle ebs", "sap", "salesforce", "servicenow",
    "full stack", "devops", "agile", "scrum", "kanban",
    "test driven development", "tdd", "bdd",
}

SKILL_SYNONYMS: dict[str, str] = {
    "ml": "machine learning", "dl": "deep learning",
    "nlp": "natural language processing", "cv": "computer vision",
    "torch": "pytorch", "sklearn": "scikit-learn",
    "scikit learn": "scikit-learn", "postgres": "postgresql",
    "k8s": "kubernetes", "tf": "tensorflow",
    "js": "javascript", "ts": "typescript",
    "node": "node.js", "nodejs": "node.js",
    "react.js": "react", "reactjs": "react",
    "mongo": "mongodb", "llms": "llm",
    "large language models": "llm", "gen ai": "generative ai",
    "genai": "generative ai", "rag": "retrieval augmented generation",
}

_SKILL_BLACKLIST = [
    re.compile(p, re.IGNORECASE) for p in [
        r"^\d+\s*%", r"\d{4}",
        r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b",
        r"\b(present|current|till|since|ongoing)\b",
        r"^(hobbies?|languages?|education|experience|projects?|"
        r"certifications?|awards?|skills?|summary|profile|objective|"
        r"references?|activities?)$",
        r"\b(travelling|cooking|reading|gaming|photography|swimming|yoga)\b",
        r"\b(cgpa|gpa|percentage|marks|score|grade|rank)\b",
        r"^[\s•·\-–—]\s*$",
        r"[<>{}()\[\]]",
    ]
]


def _is_valid_skill(s: str) -> bool:
    s = s.strip()
    if not s or len(s) < 2 or len(s) > 60:
        return False
    for pat in _SKILL_BLACKLIST:
        if pat.search(s):
            return False
    return True


def extract_skills(sections: dict, raw_text: str) -> list[str]:
    found: set[str] = set()

    # 1. Skills section (highest precision)
    skills_text = _get_section_text(sections, SectionType.SKILLS)
    if skills_text.strip():
        for line in skills_text.splitlines():
            for part in re.split(r"[,|•·;/\t]", line):
                clean = re.sub(r"^[\s•·\-–—:]+|[\s•·\-–—:]+$", "", part).strip()
                if _is_valid_skill(clean) and 2 <= len(clean) <= 50:
                    found.add(clean)

    # 2. Vocabulary scan
    lower = raw_text.lower()
    for skill in TECH_SKILLS:
        if re.search(r"\b" + re.escape(skill) + r"\b", lower):
            syn = SKILL_SYNONYMS.get(skill, skill)
            display = syn.upper() if len(syn) <= 3 else (
                syn.title() if len(syn) <= 5 else syn.capitalize()
            )
            found.add(display)

    return sorted(s for s in found if _is_valid_skill(s))


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 9 — PROJECTS
# ═══════════════════════════════════════════════════════════════════════════════

def extract_projects(sections: dict) -> list[str]:
    proj_text = _get_section_text(sections, SectionType.PROJECTS)
    if not proj_text.strip():
        return []

    projects: list[str] = []
    current_title: Optional[str] = None

    for line in proj_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        clean = re.sub(r"^[\s•·\-–—▪◦\uf0b7\*]+", "", stripped).strip()

        if "|" in clean:
            if current_title:
                projects.append(current_title)
            current_title = clean.split("|")[0].strip()
        elif (
            5 < len(clean) < 80
            and not re.search(
                r"\b(developed|built|designed|implemented|used|created|"
                r"responsible|managed|integrated|deployed)\b",
                clean, re.IGNORECASE,
            )
            and not _DEGREE_RE.search(clean)
            and not _DATE_RANGE_RE.search(clean)
            and (clean[0].isupper() if clean else False)
        ):
            if not current_title:
                current_title = clean

        if len(projects) >= 10:
            break

    if current_title and current_title not in projects and len(projects) < 10:
        projects.append(current_title)

    return list(dict.fromkeys(projects))[:10]


# ═══════════════════════════════════════════════════════════════════════════════
#  STEP 10 — POST-PROCESSING VALIDATION
# ═══════════════════════════════════════════════════════════════════════════════

def _safe_lower(val) -> str:
    """Return val.lower() safely — returns '' if val is None."""
    return (val or "").lower()


def _validate_and_fix(parsed: dict) -> dict:
    # e.get("display") can return None if the key exists but value is None.
    # Use _safe_lower() everywhere instead of .get("key","").lower()
    edu_display_set = {
        _safe_lower(e.get("display")) for e in parsed.get("education", [])
    }

    # Certifications: remove education entries that slipped in
    parsed["certifications"] = [
        c for c in parsed.get("certifications", [])
        if c is not None
        and _safe_lower(c) not in edu_display_set
        and not _CERT_ANTI_PATTERNS.search(c)
    ]

    # Companies: remove educational institutions
    parsed["previous_companies"] = [
        c for c in parsed.get("previous_companies", [])
        if c is not None
        and _is_valid_company(c)
        and not any(
            _safe_lower(c) in _safe_lower(e.get("display"))
            for e in parsed.get("education", [])
        )
    ]

    # Education: dedup — institution and degree can be None, guard both
    seen_edu: set[tuple] = set()
    clean_edu: list[dict] = []
    for edu in parsed.get("education", []):
        key = (
            _safe_lower(edu.get("degree")),
            _safe_lower(edu.get("institution")),
        )
        if key not in seen_edu:
            clean_edu.append(edu)
            seen_edu.add(key)
    parsed["education"] = clean_edu

    # Skills: filter None then case-insensitive dedup
    seen_skills: set[str] = set()
    clean_skills: list[str] = []
    for s in parsed.get("skills", []):
        if s is None:
            continue
        k = _safe_lower(s)
        if k not in seen_skills:
            clean_skills.append(s)
            seen_skills.add(k)
    parsed["skills"] = clean_skills

    return parsed


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

def parse_resume(path: Path) -> dict:
    """
    Full ATS-grade resume parsing:
      1. Extract text
      2. Detect section boundaries
      3. Extract each field from its correct section only
      4. Validate + cross-check + normalize
    """
    try:
        raw_text = extract_text(path)
    except Exception as exc:
        logger.error(f"Text extraction failed for {path.name}: {exc}")
        return {"raw_text": "", "file_name": path.name, "file_path": str(path)}

    if not raw_text.strip():
        logger.warning(f"Empty text from {path.name}")
        return {"raw_text": "", "file_name": path.name, "file_path": str(path)}

    sections = parse_sections(raw_text)
    detected = list(sections.keys())
    logger.debug(f"Sections in {path.name}: {detected}")

    nlp = get_nlp()
    edu_list = extract_education(sections)

    result: dict = {
        "name":               extract_name(raw_text, nlp),
        "email":              extract_email(raw_text),
        "phone":              extract_phone(raw_text),
        "skills":             extract_skills(sections, raw_text),
        "education":          edu_list,
        "education_display":  education_display(edu_list),
        "experience_years":   extract_experience_years(raw_text),
        "previous_companies": extract_companies(sections, nlp),
        "certifications":     extract_certifications(sections),
        "projects":           extract_projects(sections),
        "raw_text":           raw_text,
        "file_name":          path.name,
        "file_path":          str(path),
    }

    result = _validate_and_fix(result)

    logger.info(
        f"Parsed {path.name}: name={result['name']} | "
        f"edu={len(result['education'])} | "
        f"exp={result['experience_years']}yrs | "
        f"companies={len(result['previous_companies'])} | "
        f"certs={len(result['certifications'])} | "
        f"skills={len(result['skills'])}"
    )
    return result